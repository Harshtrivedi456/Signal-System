import speech_recognition as sr
from transformers import MarianMTModel, MarianTokenizer
import tkinter as tk
from tkinter import ttk  # For the dropdown menu (ComboBox)
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

# Initialize the speech recognizer
recognizer = sr.Recognizer()

# Load translation model and tokenizer (to be updated dynamically)
language_models = {
    "English": 'Helsinki-NLP/opus-mt-en-fr',
    "Gujarati": 'Helsinki-NLP/opus-mt-gu-en',
    "French": 'Helsinki-NLP/opus-mt-fr-en',
    "Spanish": 'Helsinki-NLP/opus-mt-es-en',
    "Hindi": 'Helsinki-NLP/opus-mt-hi-en',
    "Telugu": 'Helsinki-NLP/opus-mt-te-en'
}

# Create a Word document to save translations
doc = Document()
doc.add_heading('Real-time Translated Subtitles', 0)

# Dictionary to map target language names to their model codes for translation
target_language_codes = {
    "English": 'Helsinki-NLP/opus-mt-en-te',
    "Gujarati": 'Helsinki-NLP/opus-mt-en-gu',
    "French": 'Helsinki-NLP/opus-mt-en-fr',
    "Spanish": 'Helsinki-NLP/opus-mt-en-es',
    "Hindi": 'Helsinki-NLP/opus-mt-en-hi',
    "Telugu": 'Helsinki-NLP/opus-mt-en-te'
}

# Function to load models dynamically based on selected source language
def update_source_model(selected_language):
    global source_tokenizer, source_model
    model_name = language_models[selected_language]
    source_tokenizer = MarianTokenizer.from_pretrained(model_name)
    source_model = MarianMTModel.from_pretrained(model_name)
    print(f"Source language model updated to: {model_name}")

# Function to save translated text to a Word document
def save_to_word(translated_text):
    doc.add_paragraph(translated_text)
    doc.save('translated_subtitles.docx')

# Function to translate text to target language
def translate(text):
    # Tokenize the input text
    inputs = source_tokenizer(text, return_tensors="pt", padding=True)
    
    # Generate translated text
    translated = source_model.generate(**inputs)
    
    # Decode the generated tokens to text
    translated_text = source_tokenizer.decode(translated[0], skip_special_tokens=True)
    return translated_text

# Function to continuously listen for audio input
def continuous_listen():
    with sr.Microphone() as source:
        recognizer.adjust_for_ambient_noise(source)
        
        while True:
            try:
                print("Listening...")
                audio = recognizer.listen(source, timeout=15 , phrase_time_limit=15)
                
                # Recognize the speech using Google Web Speech API
                source_language_code = source_language_var.get().lower()
                text = recognizer.recognize_google(audio, language=source_language_code)
                print(f"Recognized ({source_language_var.get()}): {text}")
                
                # Translate the recognized text to the selected target language
                target_language_code = target_language_var.get()
                translated_text = translate(text)
                print(f"Translated ({target_language_code}): {translated_text}")
                
                # Update the subtitle label in the GUI
                subtitle_label.config(text=translated_text)
                root.update_idletasks()

                # Save the translated text to the Word document
                save_to_word(translated_text)
                
                if text.lower() == "exit":
                    print("Exiting...")
                    break

            except sr.UnknownValueError:
                print("Sorry, I did not understand that.")
                subtitle_label.config(text="Sorry, I did not understand that.")
                root.update_idletasks()

            except sr.RequestError as e:
                print(f"Could not request results; {e}")
                subtitle_label.config(text=f"Error: {e}")
                root.update_idletasks()

            except KeyboardInterrupt:
                print("Process interrupted by user.")
                break

# Function to send an email with the translated document to multiple recipients
def send_email(recipients, document_path):
    sender_email = "youremail@example.com"
    sender_password = "yourpassword"
    
    # Loop through the list of recipients
    for recipient_email in recipients:
        # Create the email message
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = recipient_email
        msg['Subject'] = "Translated Subtitles Document"
        
        # Add the body to the email
        body = "Attached is the translated subtitles document from today's session."
        msg.attach(MIMEText(body, 'plain'))
        
        # Attach the Word document
        with open(document_path, 'rb') as file:
            part = MIMEApplication(file.read(), Name='translated_subtitles.docx')
            part['Content-Disposition'] = 'attachment; filename="translated_subtitles.docx"'
            msg.attach(part)
        
        # Set up the SMTP server and send the email
        try:
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, recipient_email, msg.as_string())
            server.quit()
            print(f"Email sent to {recipient_email}!")
        except Exception as e:
            print(f"Failed to send email to {recipient_email}: {e}")

# Function to start listening
def start_listening():
    continuous_listen()

# Function to end the session and send the document via email
def end_session():
    save_to_word("Session ended.")
    
    # List of student email addresses
    recipients = [
        "student1@example.com",
        "student2@example.com",
        # Add the remaining 25 email addresses
    ]
    
    # Send email to all recipients
    send_email(recipients, 'translated_subtitles.docx')
    print("Session ended and document sent to all students.")

# Initialize the GUI window
root = tk.Tk()
root.title("Real-time Translated Subtitles")
root.geometry("800x250") 
root.attributes('-topmost', True)

# Create a label to display translated subtitles
subtitle_label = tk.Label(root, text="Subtitles will appear here...", font=("Helvetica", 20), wraplength=700, justify="center")
subtitle_label.pack(side='bottom', pady=10)

# Dropdown menu for source language selection
source_language_label = tk.Label(root, text="Select Source Language:", font=("Helvetica", 12))
source_language_label.pack(pady=5)

source_language_var = tk.StringVar()
source_language_dropdown = ttk.Combobox(root, textvariable=source_language_var, values=list(language_models.keys()), font=("Helvetica", 12))
source_language_dropdown.current(0)  # Default selection
source_language_dropdown.pack(pady=5)

# Dropdown menu for target language selection
target_language_label = tk.Label(root, text="Select Target Language:", font=("Helvetica", 12))
target_language_label.pack(pady=5)

target_language_var = tk.StringVar()
target_language_dropdown = ttk.Combobox(root, textvariable=target_language_var, values=list(target_language_codes.keys()), font=("Helvetica", 12))
target_language_dropdown.current(0)  # Default selection
target_language_dropdown.pack(pady=5)

# Function to handle source language selection change
def on_source_language_change(event):
    selected_language = source_language_var.get()
    update_source_model(selected_language)

source_language_dropdown.bind("<<ComboboxSelected>>", on_source_language_change)

# Start listening for audio input after 1 second
root.after(1000, start_listening)

# Add a button to end the session and send the email to all students
end_button = tk.Button(root, text="End Session and Send Email to All", command=end_session)
end_button.pack(side='top', pady=10)

root.mainloop()
