import tkinter as tk
from tkinter import ttk
import speech_recognition as sr
from transformers import MarianMTModel, MarianTokenizer
from docx import Document
# import smtplib8
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMETex
from email.mime.application import MIMEApplication

# Initialize the speech recognizer
recognizer = sr.Recognizer()

# Hindi language model setup for English to Hindi translation
language_model_name = 'Helsinki-NLP/opus-mt-en-hi'
tokenizer = MarianTokenizer.from_pretrained(language_model_name)
model = MarianMTModel.from_pretrained(language_model_name)

# Initialize Word document
doc = Document()
doc.add_heading('Real-time Translated Subtitles', 0)

# Function to save translated text to a Word document
def save_to_word(translated_text):
    doc.add_paragraph(translated_text)
    doc.save('translated_subtitles.docx')
    print("Saved to document.")

# Function to translate text from English to Hindi
def translate(text):
    inputs = tokenizer(text, return_tensors="pt", padding=True)
    translated = model.generate(**inputs)
    translated_text = tokenizer.decode(translated[0], skip_special_tokens=True)
    return translated_text

# Function to continuously listen for audio input
def continuous_listen():
    with sr.Microphone() as source:
        recognizer.adjust_for_ambient_noise(source)
        while True:
            try:
                print("Listening...")
                audio = recognizer.listen(source, timeout=15, phrase_time_limit=15)
                
                text = recognizer.recognize_google(audio, language='en-IN')  # Recognize English speech
                print(f"Recognized (English): {text}")

                # Translate the recognized English text to Hindi
                translated_text = translate(text)
                print(f"Translated: {translated_text}")

                # Update the subtitle label in the GUI
                subtitle_label.config(text=translated_text)
                root.update_idletasks()

                # Save the translated text to the Word document
                save_to_word(translated_text)

                if text.lower() == "exit":
                    print("Exiting...")
                    break

            except sr.UnknownValueError:
                print("Sorry, I did not understand that.")
                subtitle_label.config(text="क्षमा करें, समझ में नहीं आया।")
                root.update_idletasks()

            except sr.RequestError as e:
                print(f"Could not request results; {e}")
                subtitle_label.config(text=f"त्रुटि: {e}")
                root.update_idletasks()

            except KeyboardInterrupt:
                print("Process interrupted by user.")
                break

# Function to send email with the translated document
def send_email(recipients, document_path):
    sender_email = "youremail@example.com"
    sender_password = "yourpassword"

    for recipient_email in recipients:
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = recipient_email
        msg['Subject'] = "Translated Subtitles Document"

        body = "निम्नलिखित दस्तावेज़ इस सत्र के अनुवादित सबटाइटल से संबंधित है।"
        msg.attach(MIMEText(body, 'plain'))

        with open(document_path, 'rb') as file:
            part = MIMEApplication(file.read(), Name='translated_subtitles.docx')
            part['Content-Disposition'] = 'attachment; filename="translated_subtitles.docx"'
            msg.attach(part)

        try:
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, recipient_email, msg.as_string())
            server.quit()
            print(f"Email sent to {recipient_email}!")
        except Exception as e:
            print(f"Failed to send email to {recipient_email}: {e}")

# Function to start listening
def start_listening():
    continuous_listen()

# Function to end session and send email
def end_session():
    save_to_word("Session ended.")
    recipients = ["student1@example.com", "student2@example.com"]  # Add more recipients
    send_email(recipients, 'translated_subtitles.docx')
    print("Session ended and document sent to all students.")

# GUI setup
root = tk.Tk()
root.title("Real-time Translated Subtitles")
root.geometry("800x250")
root.attributes('-topmost', True)

# Subtitle label
subtitle_label = tk.Label(root, text="सबटाइटल यहां दिखाई देंगे...", font=("Helvetica", 20), wraplength=700, justify="center")
subtitle_label.pack(side='bottom', pady=10)

# Start listening after 1 second
root.after(1000, start_listening)
root.mainloop()

