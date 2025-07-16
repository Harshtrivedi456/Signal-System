import speech_recognition as sr
from transformers import MarianMTModel, MarianTokenizer
import tkinter as tk
from tkinter import ttk
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

# Initialize the speech recognizer
recognizer = sr.Recognizer()
    
# Load translation models for different languages
language_models = {
    "English": 'Helsinki-NLP/opus-mt-en-fr',
    "Gujarati": 'Helsinki-NLP/opus-mt-gu-en',
    #  "French": 'Helsinki-NLP/opus-mt-fr-en',
    #  "Spanish": 'Helsinki-NLP/opus-mt-es-en',
    "Hindi": 'Helsinki-NLP/opus-mt-hi-en',
    "Telugu": 'Helsinki-NLP/opus-mt-te-en'
}

# Target language codes for translation
target_language_codes = {
    "English": 'Helsinki-NLP/opus-mt-en-te',
    "Gujarati": 'Helsinki-NLP/opus-mt-en-gu',
     "French": 'Helsinki-NLP/opus-mt-en-fr',
     "Spanish": 'Helsinki-NLP/opus-mt-en-es',
    "Hindi": 'Helsinki-NLP/opus-mt-en-hi',
    "Telugu": 'Helsinki-NLP/opus-mt-en-te'
}

# Initialize Word document
doc = Document()
doc.add_heading('Real-time Translated Subtitles', 0)

# Default source language model
default_language = "English"
default_model_name = language_models[default_language]
source_tokenizer = MarianTokenizer.from_pretrained(default_model_name)
source_model = MarianMTModel.from_pretrained(default_model_name)

# Function to update source language model
def update_source_model(selected_language):
    global source_tokenizer, source_model
    model_name = language_models[selected_language]
    source_tokenizer = MarianTokenizer.from_pretrained(model_name)
    source_model = MarianMTModel.from_pretrained(model_name)
    print(f"Source language model updated to: {model_name}")


# Function to save translated text to a Word document
def save_to_word(translated_text):
    doc.add_paragraph(translated_text)
    doc.save('translated_subtitles.docx')

# Function to translate text
def translate(text):
    global source_tokenizer, source_model
    if source_tokenizer is None or source_model is None:
        raise RuntimeError("Source language model is not initialized. Please select a source language.")
    
    # Tokenize the input text
    inputs = source_tokenizer(text, return_tensors="pt", padding=True)
    
    # Generate translated text
    translated = source_model.generate(**inputs)
    
    # Decode the generated tokens to text
    translated_text = source_tokenizer.decode(translated[0], skip_special_tokens=True)
    return translated_text

# Function to continuously listen for audio input
import time

def continuous_listen():
    with sr.Microphone() as source:
        recognizer.adjust_for_ambient_noise(source)
        while True:
            try:
                print("Listening...")
                audio = recognizer.listen(source, timeout=15, phrase_time_limit=15)

                source_language_code = source_language_var.get().lower()
                print(f"Source Language: {source_language_code}")  # Debug line

                text = recognizer.recognize_google(audio, language=source_language_code)
                print(f"Recognized ({source_language_var.get()}): {text}")

                # Translate the recognized text
                translated_text = translate(text)
                print(f"Translated: {translated_text}")

                # Update the subtitle label in the GUI
                subtitle_label.config(text=translated_text)
                root.update_idletasks()

                # Save the translated text to the Word document
                save_to_word(translated_text)

                if text.lower() == "exit":
                    print("Exiting...")
                    break

            except sr.UnknownValueError:
                print("Sorry, I did not understand that.")
                subtitle_label.config(text="Sorry, I did not understand that.")
                root.update_idletasks()

            except sr.RequestError as e:
                print(f"Could not request results; {e}")
                subtitle_label.config(text=f"Error: {e}")
                root.update_idletasks()

            except KeyboardInterrupt:
                print("Process interrupted by user.")
                break

            except TimeoutError:
                print("No audio input detected. Retrying...")
                continue  # Retry listening

            except Exception as e:
                print(f"An unexpected error occurred: {e}")
                subtitle_label.config(text="Error occurred. Restarting...")
                root.update_idletasks()
                break


# Function to send email with the translated document
def send_email(recipients, document_path):
    sender_email = "youremail@example.com"
    sender_password = "yourpassword"

    for recipient_email in recipients:
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = recipient_email
        msg['Subject'] = "Translated Subtitles Document"

        body = "Attached is the translated subtitles document from today's session."
        msg.attach(MIMEText(body, 'plain'))

        with open(document_path, 'rb') as file:
            part = MIMEApplication(file.read(), Name='translated_subtitles.docx')
            part['Content-Disposition'] = 'attachment; filename="translated_subtitles.docx"'
            msg.attach(part)

        try:
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, recipient_email, msg.as_string())
            server.quit()
            print(f"Email sent to {recipient_email}!")
        except Exception as e:
            print(f"Failed to send email to {recipient_email}: {e}")

# Function to start listening
def start_listening():
    continuous_listen()

# Function to end session and send email
def end_session():
    save_to_word("Session ended.")
    recipients = ["student1@example.com", "student2@example.com"]  # Add more recipients
    send_email(recipients, 'translated_subtitles.docx')
    print("Session ended and document sent to all students.")

# GUI setup
root = tk.Tk()
root.title("Real-time Translated Subtitles")
root.geometry("800x250")
root.attributes('-topmost', True)

# Subtitle label
subtitle_label = tk.Label(root, text="Subtitles will appear here...", font=("Helvetica", 20), wraplength=700, justify="center")
subtitle_label.pack(side='bottom', pady=10)

# Source language dropdown
source_language_var = tk.StringVar(value="English")
source_language_dropdown = ttk.Combobox(root, textvariable=source_language_var, values=list(language_models.keys()), font=("Helvetica", 12))
source_language_dropdown.current(0)
source_language_dropdown.pack(pady=5)

# Target language dropdown
target_language_var = tk.StringVar(value="English")
target_language_dropdown = ttk.Combobox(root, textvariable=target_language_var, values=list(target_language_codes.keys()), font=("Helvetica", 12))
target_language_dropdown.current(0)
target_language_dropdown.pack(pady=5)

# Update source language model on selection
def on_source_language_change(event):
    selected_language = source_language_var.get()
def on_source_language_change(event):
    selected_language = source_language_var.get()
    update_source_model(selected_language)

source_language_dropdown.bind("<<ComboboxSelected>>", on_source_language_change)


# End session button
end_button = tk.Button(root, text="End Session and Send Email to All", command=end_session)
end_button.pack(side='top', pady=10)

# Start listening after 1 second
root.after(1000, start_listening)
root.mainloop()
